"""import pdfplumber
from docx import Document
from app.utils.text_cleaner import clean_text


def parse_pdf_jd(file_path: str) -> str:
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return clean_text(text)


def parse_docx_jd(file_path: str) -> str:
    doc = Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return clean_text(text)


def parse_jd(file_path: str) -> str:
    if file_path.endswith(".pdf"):
        return parse_pdf_jd(file_path)
    elif file_path.endswith(".docx"):
        return parse_docx_jd(file_path)
    else:
        raise ValueError("Unsupported JD format")
"""

import pdfplumber
from docx import Document

def parse_jd(file_path: str) -> str:
    if file_path.endswith(".pdf"):
        try:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text.strip()
        except Exception:
            return ""

    elif file_path.endswith(".docx"):
        try:
            doc = Document(file_path)
            return "\n".join(p.text for p in doc.paragraphs if p.text)
        except Exception:
            return ""

    return ""
